import os
from pathlib import Path
from typing import List
import fitz
from PIL import Image
import docx
from app.services.ocr_service import perform_ocr

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF using PyMuPDF.
    Falls back to OCR for scanned PDFs with no text layer.
    
    Args:
        file_path: Path to PDF file
    
    Returns:
        Extracted text content
    """
    text_parts = []
    
    try:
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                text_parts.append(text)
            else:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                
                temp_img_path = f"{file_path}_page_{page_num}.png"
                img.save(temp_img_path)
                
                ocr_text = perform_ocr(temp_img_path)
                text_parts.append(ocr_text)
                
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
        
        doc.close()
        
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"
    
    return "\n".join(text_parts)

def extract_text_from_word(file_path: str) -> str:
    """
    Extract text from Word document (.docx).
    
    Args:
        file_path: Path to Word file
    
    Returns:
        Extracted text content
    """
    try:
        doc = docx.Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        return f"Error extracting Word document: {str(e)}"

def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from plain text file.
    
    Args:
        file_path: Path to text file
    
    Returns:
        File content
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()
        except Exception as e:
            return f"Error reading text file: {str(e)}"
    except Exception as e:
        return f"Error reading text file: {str(e)}"

def extract_text_from_file(file_path: str, file_type: str) -> str:
    """
    Extract text from uploaded file based on type.
    
    Supports:
    - PDF: PyMuPDF with OCR fallback for scanned pages
    - Word (.docx): python-docx
    - Images (jpg, jpeg, png): Tesseract OCR
    - Text (.txt): Direct reading
    
    Args:
        file_path: Path to the uploaded file
        file_type: File extension (pdf, jpg, png, docx, txt)
    
    Returns:
        Extracted text content
    """
    
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    
    elif file_type == "docx":
        return extract_text_from_word(file_path)
    
    elif file_type in ["jpg", "jpeg", "png"]:
        return perform_ocr(file_path)
    
    elif file_type == "txt":
        return extract_text_from_txt(file_path)
    
    else:
        return f"Unsupported file type: {file_type}"
